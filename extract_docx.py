#!/usr/bin/env python3
"""
Script to extract text content from a DOCX file
"""
import sys
from docx import Document

def extract_docx_content(file_path):
    try:
        doc = Document(file_path)

        # Extract all paragraphs
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())

        # Extract tables
        for table in doc.tables:
            content.append("\n--- TABLE ---")
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_content.append(cell_text)
                if row_content:
                    content.append(" | ".join(row_content))
            content.append("--- END TABLE ---\n")

        return "\n".join(content)

    except Exception as e:
        return f"Error extracting content: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <path_to_docx_file>")
        sys.exit(1)

    file_path = sys.argv[1]
    content = extract_docx_content(file_path)
    print(content)