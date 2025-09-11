#!/usr/bin/env python3
"""
Q-Gen Questionnaire Generator (.docx) — brand + color formatting

Adds:
  • Logo at top (looks for /ui/icons/logo.png; env overrides)
  • Color-codes flags: [exclusive] BLUE, [terminate] RED
  • Cleaner table/header styling

Usage:
  pip install python-docx
  python generate_questionnaire.py --project project.json --out questionnaire.docx
  # optional:
  #   --hide-codes  (don’t show numeric codes for options)
  #   --logo PATH   (override logo path)
"""

import argparse, json, re, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------- helpers ----------

BLUE = RGBColor(0x00, 0x5E, 0xFF)   # “exclusive”
RED  = RGBColor(0xCC, 0x00, 0x00)   # “terminate”
GREY = RGBColor(0x55, 0x55, 0x55)   # meta text like [Routing]

def is_screener(qid: str) -> bool:
    return str(qid or "").strip().upper().startswith("S")

def q_sort_key(qid: str):
    qid = str(qid or "")
    group = 0 if is_screener(qid) else 1
    m = re.match(r"^[SQ](\d+)", qid.upper())
    num = int(m.group(1)) if m else 10_000
    return (group, num, qid)

def add_para(doc: Document, text: str, *, bold=False, italic=False, size=11, after=6, align=None, color: RGBColor|None=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    for r in p.runs:
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p

def add_matrix(doc: Document, headers, rows, *, col_widths=None):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for j, h in enumerate(headers):
        c = tbl.cell(0, j)
        c.text = str(h)
        # bold header text
        if c.paragraphs and c.paragraphs[0].runs:
            for r in c.paragraphs[0].runs:
                r.bold = True

    # body
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            tbl.cell(i, j).text = str(val)

    # widths (best effort)
    if col_widths:
        try:
            for j, w in enumerate(col_widths):
                for i in range(len(rows)+1):
                    tbl.cell(i, j).width = Inches(w)
        except Exception:
            pass

    add_para(doc, "", after=8)
    return tbl

def badges_for(q):
    qid = (q.get("id") or "").upper()
    notes = (q.get("notes") or "").lower()
    tags = []
    if "_H" in qid or "hidden" in notes:
        tags.append("Hidden")
    if "_R" in qid or "red herring" in notes or "qc" in notes:
        tags.append("Red herring/QC")
    return tags

def render_option_line(p, prefix_text: str, label: str, flags):
    """
    Writes one option line into paragraph p with colored flag runs.
      prefix_text -> regular (e.g., "1. " or "• ")
      label       -> regular
      [exclusive] -> BLUE
      [terminate] -> RED
    """
    r1 = p.add_run(prefix_text + label)
    r1.font.size = Pt(11)

    blk = []
    if flags:
        # Order: exclusive (blue), terminate (red)
        if "exclusive" in flags: blk.append((" [exclusive]", BLUE))
        if "terminate" in flags: blk.append((" [terminate]", RED))
    for txt, col in blk:
        rr = p.add_run(txt)
        rr.font.size = Pt(11)
        rr.font.color.rgb = col

def render_options_list(doc: Document, options, show_codes=True):
    if not options:
        return
    for opt in options:
        code = opt.get("code", "")
        label = opt.get("label", "") or ""
        flags = set()
        if opt.get("exclusive"): flags.add("exclusive")
        if opt.get("terminate"): flags.add("terminate")

        p = doc.add_paragraph()
        if show_codes and code not in ("", None):
            render_option_line(p, f"{code}. ", label, flags)
        else:
            render_option_line(p, "• ", label, flags)
        p.paragraph_format.space_after = Pt(2)
    add_para(doc, "", after=6)

def logo_path_override(cli_logo: str|None) -> Path|None:
    """
    Search order for logo:
      1) --logo argument
      2) QGEN_LOGO env var
      3) ../ui/icons/logo.png (relative to this file)
      4) ./ui/icons/logo.png (if running from project root)
    """
    # CLI
    if cli_logo:
        p = Path(cli_logo)
        return p if p.exists() else None
    # ENV
    env = os.getenv("QGEN_LOGO")
    if env and Path(env).exists():
        return Path(env)
    # ../ui/icons/logo.png (backend → ui)
    here = Path(__file__).resolve()
    candidate = here.parent.parent / "ui" / "icons" / "logo.png"
    if candidate.exists():
        return candidate
    # ./ui/icons/logo.png (project root)
    candidate = Path.cwd() / "ui" / "icons" / "logo.png"
    if candidate.exists():
        return candidate
    return None

# ---------- main ----------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--project", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--hide-codes", action="store_true")
    ap.add_argument("--logo", help="Path to logo PNG to show in header")
    args = ap.parse_args()

    with open(args.project, "r", encoding="utf-8") as f:
        data = json.load(f)

    proj = data.get("project", {})
    questions = data.get("questions", [])
    questions.sort(key=lambda q: q_sort_key(q.get("id","")))

    doc = Document()

    # Logo + Title
    lp = logo_path_override(args.logo)
    if lp:
        try:
            doc.add_picture(str(lp), width=Inches(1.5))
            add_para(doc, "", after=4)
        except Exception:
            # don't fail hard if PIL support is missing for PNG variants
            pass

    title = proj.get("name") or "Survey Questionnaire"
    client = proj.get("client") or ""
    add_para(doc, title, bold=True, size=16, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    if client:
        add_para(doc, f"Client: {client}", italic=True, size=10, after=12, color=GREY)

    current_section = None
    for q in questions:
        qid = q.get("id","")
        qtext = (q.get("text") or "").strip()
        qtype = (q.get("type") or "").strip()
        special = q.get("special_verbiage") or ""
        routing = q.get("routing") or ""
        base = q.get("base") or {}
        base_verb = base.get("verbiage") or ""
        base_def = base.get("definition") or ""
        options = q.get("options") or []
        scale = q.get("scale") or {}
        labels = [l for l in (scale.get("labels") or []) if str(l).strip()]
        statements = [s for s in (q.get("statements") or []) if str(s).strip()]

        # section header
        section = "Screener" if is_screener(qid) else "Main Survey"
        if current_section != section:
            add_para(doc, section, bold=True, size=14, after=6)
            current_section = section

        # question header
        header = f"{qid}. {qtext}" if qtext else f"{qid}"
        add_para(doc, header, bold=True, size=12, after=2)

        # badges
        tags = badges_for(q)
        if tags:
            add_para(doc, "[" + " | ".join(tags) + "]", italic=True, size=9, after=2, color=GREY)

        # verbiage / routing / base (render in grey meta style)
        if special:
            add_para(doc, f"[Verbiage] {special}", italic=True, size=10, after=0, color=GREY)
        if routing:
            add_para(doc, f"[Routing] {routing}", italic=True, size=10, after=0, color=GREY)
        if base_verb or base_def:
            base_line = f"[Base] {base_verb}" + (f" | Def: {base_def}" if base_def else "")
            add_para(doc, base_line, italic=True, size=10, after=4, color=GREY)

        # ---- render by type ----
        if qtype.startswith("likert"):
            # Matrix: statements x scale.labels
            if statements and labels:
                headers = ["Statement"] + labels
                rows = [[s] + [""] * len(labels) for s in statements]
                add_matrix(doc, headers, rows, col_widths=[2.8] + [1.1]*len(labels))
            else:
                if labels:
                    add_para(doc, "Response Scale: " + " | ".join(labels), italic=True, size=10, after=4, color=GREY)
                render_options_list(doc, options, show_codes=not args.hide_codes)

        elif qtype.startswith("grid_"):
            # Matrix: statements x options/scale
            if statements:
                if options:
                    opt_headers = [
                        (f"{o.get('code')}. {o.get('label')}" if not args.hide_codes and o.get('code') not in (None,"")
                         else f"{o.get('label','')}")
                        for o in options
                    ]
                    headers = ["Statement"] + opt_headers
                else:
                    # fall back to scale labels as columns
                    opt_headers = labels if labels else []
                    headers = ["Statement"] + opt_headers

                rows = [[s] + [""] * (len(headers)-1) for s in statements]
                # try to size columns sanely
                w = max(1.0, 5.5 / max(1, (len(headers)-1)))
                add_matrix(doc, headers, rows, col_widths=[2.8] + [w]*(len(headers)-1))
            else:
                render_options_list(doc, options, show_codes=not args.hide_codes)

        elif qtype in ("single", "multi", "open", "numeric", "numeric_open"):
            render_options_list(doc, options, show_codes=not args.hide_codes)
            if qtype.startswith("numeric"):
                units = (q.get("numeric") or {}).get("unit") or (q.get("numeric") or {}).get("units")
                if units:
                    add_para(doc, f"(Numeric units: {units})", italic=True, size=10, after=4, color=GREY)

        else:
            render_options_list(doc, options, show_codes=not args.hide_codes)

        add_para(doc, "", after=10)  # spacing after each question

    doc.save(args.out)
    print(f"Wrote {args.out}")

if __name__ == "__main__":
    main()
